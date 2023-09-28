import io
import re

import docx
import pandas as pd
from docx import Document

class DataProcessor:
    def validate_file(self, file):
        return True if file.endswith('.docx') else False

    def tokenize_text(self, text: str) -> list:
        """Tokenizes input text into sentences using regex"""

        regex = r"(?<!гр)(?<!см)(?<!им)(?<!\sо)(?<!\sг)(?<!\sр)(?<!\.[А-Я])(?<!\.\s[А-Я])(\. +[А-Я])(?!\.)"

        r1 = re.split(regex, text)
        sents = [r1[0] + "."]
        for i in range(1, len(r1), 2):
            sents.append(r1[i][-1] + r1[i + 1] + ".")
        sents[-1] = sents[-1][0:-1]

        sents = [sent for sent in sents if re.search("[а-яА-Яa-zA-Z]", sent)]

        return sents
    
    def tokenize_docx(self, file):
        # Your document tokenization code goes here
        # Example: Tokenize the DOCX file into sentences

        doc = Document(io.BytesIO(file))
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append(para.text.strip())

        df = pd.DataFrame(paragraphs, columns=["paragraphs"])
        df["sents"] = df["paragraphs"].apply(lambda x: self.tokenize_text(x))
        df = df.explode("sents")["sents"]
        sents = [str(sent) for sent in df.values]

        # remove rows after References
        try:
            threshold = sents.index('СПИСОК ЛИТЕРАТУРЫ')
            sents = sents[:threshold]
        except ValueError:
            pass
        
        return sents
    
    def concatenate_sentences_to_docx(self, sentences):
        """
        уже после перевода
        """
        text = '. '.join(sentences)
        document = docx.Document()
        paragraph = document.add_paragraph()
        paragraph.add_run(text)

        return document



    

